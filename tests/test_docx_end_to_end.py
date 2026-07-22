from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def test_markdown_to_docx_end_to_end(tmp_path: Path) -> None:
    source = tmp_path / "最小交底书.md"
    output = tmp_path / "最小交底书.docx"
    source.write_text(
        "# 一种测试方法\n\n## 1. 技术领域\n\n本方案涉及任务调度。\n\n"
        "| 参数 | 含义 |\n|---|---|\n| W | 窗口大小 |\n",
        encoding="utf-8",
    )
    subprocess.run(
        [sys.executable, str(ROOT / "tools" / "md_to_docx.py"), "-i", str(source), "-o", str(output)],
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    assert output.stat().st_size > 1_000
    document = Document(output)
    assert any(paragraph.text == "一种测试方法" for paragraph in document.paragraphs)
    assert document.tables[0].cell(1, 1).text == "窗口大小"
